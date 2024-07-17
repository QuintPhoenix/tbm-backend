from docx import Document
from docx.shared import Inches
from io import BytesIO
import requests
import uuid
from htmldocx import HtmlToDocx
import markdown
from os import environ

def get_para_data(output_doc_name, paragraph):
    """
    Write the run to the new file and then set its font, bold, alignment, color etc. data.
    """

    output_para = output_doc_name.add_paragraph()
    for run in paragraph.runs:
        output_run = output_para.add_run(run.text)
        # Run's bold data
        output_run.bold = run.bold
        # Run's italic data
        output_run.italic = run.italic
        # Run's underline data
        output_run.underline = run.underline
        # Run's color data
        output_run.font.color.rgb = run.font.color.rgb
        # Run's font data
        output_run.style.name = run.style.name
    # Paragraph's alignment data
    output_para.paragraph_format.alignment = paragraph.paragraph_format.alignment


path_to_doc_folder = environ.get("GENERATED_DOC_PATH")

def generate_doc(title : str, imageLinks : list, body : list) -> str:
    doc = Document()
    doc.add_heading(title,0).alignment = 1
    
    if(imageLinks) : table = doc.add_table(rows=len(imageLinks),cols=2)
    i=0
    indexes = [[0,0],[0,1],[2,0],[2,1]]
    indexes_caption = [[1,0],[1,1],[3,0],[3,1]]
    for it in imageLinks:
        response = requests.get(it[1])
        binary_img = BytesIO(response.content)
        row = table.rows[indexes[i][0]].cells
        row[indexes[i][1]].add_paragraph().add_run().add_picture(binary_img,width=Inches(3))
        caption_row = table.rows[indexes_caption[i][0]].cells
        caption_row[indexes_caption[i][1]].text = "Credits - " + it[0]
        caption_row[indexes_caption[i][1]].vertical_alignment = 1
        i = i + 1
    
    for key in body :
        doc.add_heading(key[0])
        html_string = markdown.markdown(key[1])
        new_parser = HtmlToDocx()
        md_doc = new_parser.parse_html_string(html_string)
        for para in md_doc.paragraphs:
            get_para_data(doc,para)

    generated_doc_name = uuid.uuid4().hex + ".docx"
    doc.save(path_to_doc_folder + generated_doc_name)
    return generated_doc_name
